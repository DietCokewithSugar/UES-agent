#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
用户研究分析结论 .docx 生成辅助脚本
================================

把"按分析结论章节组织的结构化 JSON"渲染为 Word(.docx) 的通用工具，
负责两件事，避免每次分析都重写一遍 python-docx / matplotlib 样板代码：

1) 按 JSON 结构生成 .docx（标题、章节、段落、要点、结论块、表格、分页）
2) 根据 JSON 内嵌的图表数据（chart spec）自动用 matplotlib 绘制并插入，
   支持 bar / line / pie / scatter / funnel 五种类型

用法
----
    python analysis_builder.py analysis.json output.docx

analysis.json 的通用协议（blocks 块结构）见下方 ANALYSIS_EXAMPLE 与
SKILL.md Step 6 的"工具与库指引"。核心结构为：

    {
      "title": "分析结论标题",
      "subtitle": "可选副标题",
      "note": "可选：生成日期/方法简述",
      "blocks": [ ... 见 ANALYSIS_EXAMPLE ... ]
    }

排版约定（与 references/analysis_template.md 的格式规范保持一致）：
    - 正文段落（paragraph）默认段首空两格（firstLineChars=200）；
    - 核心结论用 conclusion 块：一句话结论单独一段无项目符号、
      关键数据逐条单独一段带项目符号、原因解读与结论可信度分别单独一段无项目符号；
    - 行动建议内容较多时用 bullets 逐条分段展示，不用表格。

依赖
    pip install python-docx matplotlib

说明
    - 图表渲染需写入文件，默认输出到 <output> 同目录的临时 PNG，插入后自动清理。
    - 若系统缺少中文字体，图内中文可能显示为方块；脚本自动尝试微软雅黑等常见中文字体。
"""

import json
import os
import sys

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
except ImportError:
    sys.exit("缺少依赖 python-docx，请先执行: pip install python-docx")

# --------------------------------------------------------------------------- #
# 全文排版规范：字体、字号、行距、段前段后
# --------------------------------------------------------------------------- #
CN_FONT = "微软雅黑"        # 全文统一字体（同时覆盖 ASCII / 东亚文字）
BODY_SIZE = 11              # 正文字号（pt）
LINE_SPACING = 1.15         # 全文固定行距（倍）
TITLE_SIZE = 18             # 文档主标题字号（pt）
NOTE_SIZE = 9               # 注释 / 图注字号（pt）
HEADING_SPEC = {            # 分级标题: 字号 / 段前 / 段后（pt）
    1: {"size": 16, "before": 12, "after": 6},
    2: {"size": 14, "before": 10, "after": 4},
    3: {"size": 12, "before": 8, "after": 4},
}


def _set_font(run, size=None, bold=None, color=None):
    """为 run 设置中英文字体（微软雅黑）。"""
    run.font.name = CN_FONT
    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), CN_FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def _set_style_font(style):
    """将样式（Normal / Heading 等）的默认字体设为微软雅黑。"""
    style.font.name = CN_FONT
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), CN_FONT)


def _set_first_line_indent(paragraph, chars=2.0, font_size=BODY_SIZE):
    """设置段首缩进（按字符数，中文排版规范=首行空两格）。

    字符缩进会随字号缩放（w:firstLineChars=chars*100）；
    同时写入 firstLine（twips）作为不识别 firstLineChars 的编辑器的兼容值。
    """
    pPr = paragraph._p.get_or_add_pPr()
    ind = pPr.find(qn("w:ind"))
    if ind is None:
        ind = pPr.makeelement(qn("w:ind"), {})
        pPr.append(ind)
    ind.set(qn("w:firstLineChars"), str(int(chars * 100)))
    ind.set(qn("w:firstLine"), str(int(chars * font_size * 20)))


def _apply_doc_fonts(doc):
    """对整个文档的内置样式应用排版规范：字体 + 分级字号 + 段前段后。"""
    style_specs = {
        "Normal": {"size": BODY_SIZE, "before": 0, "after": 6},
        "Title": {"size": TITLE_SIZE, "before": 0, "after": 12},
        "Subtitle": {"size": 12, "before": 0, "after": 6},
        "Caption": {"size": NOTE_SIZE, "before": 2, "after": 2},
    }
    for level, spec in HEADING_SPEC.items():
        style_specs["Heading %d" % level] = spec
    for sname, spec in style_specs.items():
        try:
            style = doc.styles[sname]
        except KeyError:
            continue
        _set_style_font(style)
        style.font.size = Pt(spec["size"])
        style.paragraph_format.space_before = Pt(spec["before"])
        style.paragraph_format.space_after = Pt(spec["after"])
        style.paragraph_format.line_spacing = LINE_SPACING

try:
    import matplotlib
    matplotlib.use("Agg")  # 无界面后端，命令行可用
    import matplotlib.pyplot as plt
    _HAS_MPL = True
except Exception:
    _HAS_MPL = False

DARK = RGBColor(0x1F, 0x4E, 0x79)      # 深蓝（标题）
GRAY = RGBColor(0x59, 0x59, 0x59)      # 灰（注释）


# --------------------------------------------------------------------------- #
# 中文字体（避免 matplotlib 中文显示为方块）
# --------------------------------------------------------------------------- #
def _ensure_cjk_font():
    from matplotlib import font_manager as fm
    for name in ("Microsoft YaHei", "SimHei", "SimSun",
                 "Noto Sans CJK SC", "PingFang SC", "WenQuanYi Zen Hei"):
        try:
            fm.findfont(name, fallback_to_default=False)
            plt.rcParams["font.sans-serif"] = [name]
            break
        except Exception:
            continue
    plt.rcParams["axes.unicode_minus"] = False


# --------------------------------------------------------------------------- #
# 图表渲染（返回生成的 .png 路径）
# --------------------------------------------------------------------------- #
def _render_chart(spec):
    """按 spec 渲染图表。spec 必须含 'type'，其余字段因类型而异。"""
    if not _HAS_MPL:
        return None
    _ensure_cjk_font()
    ctype = spec.get("type")
    title = spec.get("title", "")
    figsize = tuple(spec.get("figsize", (8, 4)))
    out_png = spec["_out"]

    if ctype == "bar":
        fig, ax = plt.subplots(figsize=figsize)
        labels = spec["labels"]
        series = spec["series"]  # [[名称, 数值列表], ...]
        n = len(labels)
        nser = len(series)
        width = spec.get("bar_width", 0.6 / nser if nser else 0.6)
        for i, (sname, vals) in enumerate(series):
            xs = [j + i * width - width * (nser - 1) / 2 for j in range(n)]
            ax.bar(xs, vals, width=width, label=sname)
        ax.set_xticks(range(n))
        ax.set_xticklabels(labels, rotation=15, ha="right")
        ax.set_ylabel(spec.get("ylabel", ""))
        if nser > 1:
            ax.legend()
        ax.set_title(title)

    elif ctype == "line":
        fig, ax = plt.subplots(figsize=figsize)
        for sname, vals in spec["series"]:
            ax.plot(spec["x"], vals, marker="o", label=sname)
        ax.set_xlabel(spec.get("xlabel", ""))
        ax.set_ylabel(spec.get("ylabel", ""))
        if spec.get("series"):
            ax.legend()
        ax.set_title(title)

    elif ctype == "pie":
        fig, ax = plt.subplots(figsize=figsize)
        ax.pie(spec["values"], labels=spec["labels"], autopct="%1.1f%%",
               startangle=90)
        ax.axis("equal")
        ax.set_title(title)

    elif ctype == "scatter":
        # 优先级矩阵：散点 + 均值十字线（四象限）
        fig, ax = plt.subplots(figsize=figsize)
        xs = [p["x"] for p in spec["points"]]
        ys = [p["y"] for p in spec["points"]]
        labels = [p.get("label", "") for p in spec["points"]]
        ax.scatter(xs, ys, s=spec.get("size", 120))
        for x, y, la in zip(xs, ys, labels):
            ax.annotate(la, (x, y), textcoords="offset points", xytext=(6, 4))
        if spec.get("quadrant_lines", True):
            ax.axhline(sum(ys) / len(ys), color=(0.6, 0.6, 0.6),
                       linestyle="--", linewidth=1)
            ax.axvline(sum(xs) / len(xs), color=(0.6, 0.6, 0.6),
                       linestyle="--", linewidth=1)
        ax.set_xlabel(spec.get("xlabel", "影响度"))
        ax.set_ylabel(spec.get("ylabel", "置信度"))
        ax.set_title(title)

    elif ctype == "funnel":
        # 漏斗：按剩余量依次递减的横向条形
        fig, ax = plt.subplots(figsize=figsize)
        steps = spec["steps"]  # [[步骤名, 人数], ...]
        names = [s[0] for s in steps]
        vals = [s[1] for s in steps]
        ax.barh(range(len(vals)), vals)
        ax.invert_yaxis()
        ax.set_yticks(range(len(names)))
        ax.set_yticklabels(names)
        ax.set_xlabel(spec.get("xlabel", "人数"))
        ax.set_title(title)
        first = vals[0] if vals else 0
        for i, v in enumerate(vals):
            rate = "%s%%" % round(100.0 * v / first, 1) if first else ""
            ax.text(v, i, "  %s" % rate, va="center")

    elif ctype == "radar":
        # ETS 8 维度雷达图
        fig, ax = plt.subplots(figsize=figsize, subplot_kw=dict(polar=True))
        cats = spec["categories"]
        n = len(cats)
        angles = [i / float(n) * 2 * 3.14159 for i in range(n)]
        for sname, vals in spec["series"]:
            values = vals + [vals[0]]
            angs = angles + [angles[0]]
            ax.plot(angs, values, label=sname)
            ax.fill(angs, values, alpha=0.1)
        ax.set_xticks(angles)
        ax.set_xticklabels(cats)
        if spec.get("legend", True) and spec.get("series"):
            ax.legend(loc="upper right", bbox_to_anchor=(1.3, 1.15))
        ax.set_title(title)

    else:
        raise ValueError("不支持的图表类型: %s" % ctype)

    fig.tight_layout()
    fig.savefig(out_png, dpi=150)
    plt.close(fig)
    return out_png


# --------------------------------------------------------------------------- #
# docx 基础构件
# --------------------------------------------------------------------------- #
def _add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = LINE_SPACING
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(text)
    _set_font(r, size=TITLE_SIZE, bold=True, color=DARK)


def _add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    h.paragraph_format.line_spacing = LINE_SPACING
    h.paragraph_format.space_before = Pt(HEADING_SPEC[level]["before"])
    h.paragraph_format.space_after = Pt(HEADING_SPEC[level]["after"])
    for run in h.runs:
        _set_font(run, size=HEADING_SPEC[level]["size"], color=DARK)


def _add_para(doc, text, style="body", bold=False):
    p = doc.add_paragraph(text)
    p.paragraph_format.line_spacing = LINE_SPACING
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    if style == "note":
        for run in p.runs:
            _set_font(run, size=NOTE_SIZE, color=GRAY)
    else:
        # 正文段落默认段首空两格（中文排版规范）
        _set_first_line_indent(p)
        if p.runs:
            _set_font(p.runs[0], size=BODY_SIZE, bold=bold)
    return p


def _add_conclusion(doc, block):
    """核心结论块：一句话结论（加粗）/ 关键数据（逐条项目符号，无悬挂缩进）/ 结论可信度+原因解读。

    协议（conclusion 块）：
      {
        "type": "conclusion",
        "statement": "一句话结论（加粗，单独一段）",
        "data": ["关键数据1（单独一段，带项目符号、无悬挂缩进）", "关键数据2", ...],
        "interpretation": "原因解读（可选，与结论可信度合成一段，置于其后）",
        "confidence": "结论可信度：高 / 中 / 低（单独段首，与原因解读同段）"
      }
    """
    if block.get("statement"):
        _add_conclusion_statement(doc, block["statement"])
    for item in block.get("data", []):
        _add_conclusion_data(doc, item)
    # 结论可信度 + 原因解读合并为一段：结论可信度在前，原因解读在后
    tail_parts = []
    if block.get("confidence"):
        tail_parts.append(block["confidence"])
    if block.get("interpretation"):
        tail_parts.append(block["interpretation"])
    if tail_parts:
        _add_para(doc, "；".join(tail_parts))


def _add_conclusion_statement(doc, text):
    """一句话结论：加粗、单独一段、无项目符号、段首空两格。"""
    _add_para(doc, text, bold=True)


def _add_conclusion_data(doc, text):
    """关键数据：单独一段、带项目符号、无悬挂缩进（正文段首空两格内起行）。"""
    from docx.shared import Cm
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = LINE_SPACING
    pf.space_before = Pt(0)
    pf.space_after = Pt(3)
    pf.left_indent = Cm(0.74)          # 项目符号起始位置
    r1 = p.add_run("\u2022 ")
    _set_font(r1, size=BODY_SIZE)
    r2 = p.add_run(text)
    _set_font(r2, size=BODY_SIZE)


def _add_bullets(doc, items):
    for it in items:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.line_spacing = LINE_SPACING
        pf.space_before = Pt(0)
        pf.space_after = Pt(3)
        pf.left_indent = Cm(0.74)
        pf.first_line_indent = -Cm(0.74)
        r1 = p.add_run("\u2022 ")
        _set_font(r1, size=BODY_SIZE)
        r2 = p.add_run(it)
        _set_font(r2, size=BODY_SIZE)


def _add_bullets(doc, items):
    for it in items:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.line_spacing = LINE_SPACING
        pf.space_before = Pt(0)
        pf.space_after = Pt(3)
        pf.left_indent = Cm(0.74)
        pf.first_line_indent = -Cm(0.74)
        r1 = p.add_run("\u2022 ")
        _set_font(r1, size=BODY_SIZE)
        r2 = p.add_run(it)
        _set_font(r2, size=BODY_SIZE)


def _add_table(doc, headers, rows, title=None):
    if title:
        _add_heading(doc, title, 3)
    table = doc.add_table(rows=1, cols=len(headers))
    try:
        table.style = "Light Grid Accent 1"
    except KeyError:
        try:
            table.style = "Table Grid"
        except Exception:
            pass
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        for para in hdr[i].paragraphs:
            para.paragraph_format.line_spacing = LINE_SPACING
            para.paragraph_format.space_after = Pt(0)
            run = para.add_run(str(h))
            _set_font(run, size=BODY_SIZE, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            for para in cells[i].paragraphs:
                para.paragraph_format.line_spacing = LINE_SPACING
                para.paragraph_format.space_after = Pt(0)
                for run in para.runs:
                    _set_font(run, size=BODY_SIZE)


def _add_chart(doc, spec):
    if not _HAS_MPL:
        doc.add_paragraph("（图表生成失败：环境缺少 matplotlib）")
        return
    png = _render_chart(spec)
    if png and os.path.exists(png):
        img_w = Inches(float(spec.get("width", 5.2)))
        doc.add_picture(png, width=img_w)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if spec.get("caption"):
            cp = _add_para(doc, spec["caption"], "note")
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_para(doc, "", "note")
        try:
            os.remove(png)
        except OSError:
            pass


# --------------------------------------------------------------------------- #
# 主流程
# --------------------------------------------------------------------------- #
def _check_subtitle(sub):
    """副标题格式自检：只应含生成日期 + 方案类型，不得出现样本/数据内容。"""
    if not sub:
        return
    if "生成日期" not in sub:
        print("[warn] 副标题建议包含生成日期，格式：生成日期：YYYY-MM-DD ｜ 方案类型")
    forbidden = ["N=", "N ＝", "样本", "数据来源", "数据源", "问卷", "访谈",
                 "受访者", "人/", "人；", "人)", "人）", "回收", "方法："]
    hit = [w for w in forbidden if w in sub]
    if hit:
        print("[warn] 副标题不应包含样本/数据相关内容（已检出：%s）。"
              "请仅保留生成日期与方案类型，样本/数据写入正文研究概述。" % "、".join(hit))


def build_report(data, out_docx):
    doc = Document()
    _apply_doc_fonts(doc)
    _add_title(doc, data["title"])
    _check_subtitle(data.get("subtitle"))
    if data.get("subtitle"):
        p = _add_para(doc, data["subtitle"], "note")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if data.get("note"):
        _add_para(doc, data["note"], "note")
    doc.add_paragraph()

    for block in data.get("blocks", []):
        t = block.get("type")
        if t == "heading":
            _add_heading(doc, block["text"], block.get("level", 1))
        elif t == "paragraph":
            _add_para(doc, block["text"])
        elif t == "conclusion":
            _add_conclusion(doc, block)
        elif t == "bullets":
            _add_bullets(doc, block["items"])
        elif t == "table":
            _add_table(doc, block["headers"], block["rows"], block.get("title"))
        elif t == "chart":
            spec = block["spec"]
            spec["_out"] = os.path.join(os.path.dirname(os.path.abspath(out_docx)),
                                        "_tmp_chart.png")
            _add_chart(doc, spec)
        elif t == "image":
            # 直接插入已有图片（如眼动仪导出的热区图/轨迹图截图）
            img_path = block["path"]
            if os.path.exists(img_path):
                img_w = Inches(float(block.get("width", 5.2)))
                doc.add_picture(img_path, width=img_w)
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                if block.get("caption"):
                    cp = _add_para(doc, block["caption"], "note")
                    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                _add_para(doc, "（图片缺失：%s）" % img_path)
        elif t == "pagebreak":
            doc.add_page_break()
        else:
            print("[warn] 忽略未知 block type: %s" % t)

    doc.save(out_docx)
    print("已生成分析结论: %s" % out_docx)


ANALYSIS_EXAMPLE = {
    "title": "XX产品用户研究分析结论",
    "subtitle": "生成日期：2026-08-04 ｜ 方案A：客群主轴式",
    "blocks": [
        {"type": "heading", "text": "一、研究概述", "level": 1},
        {"type": "heading", "text": "1.1 研究背景与目标", "level": 2},
        {"type": "paragraph", "text": "为了解 XX 功能的体验表现，本研究采用混合方法设计。数据来源：问卷 N=156、访谈 N=12、埋点 2026-01~02（多源，结论按证据强度标注）。"},
        {"type": "bullets", "items": ["量化满意度与 NPS", "挖掘使用中的深层痛点"]},
        {"type": "heading", "text": "2.1 注册流程门槛较高", "level": 2},
        {"type": "conclusion",
         "statement": "注册流程是当前满意度最低、流失最高的关键环节，建议优先优化。",
         "data": ["该功能满意度均值 2.8/5，显著低于登录（3.5/5）与转账（3.9/5）（问卷 N=156）",
                  "注册至首次转账的转化率仅 52%，漏斗首步流失最大（埋点 2026-01~02）",
                  "\"注册要填的信息太多，好几次想放弃\"（受访者 P03）"],
         "interpretation": "本研究认为，高流失主要源于信息采集门槛，而非功能故障。",
         "confidence": "结论可信度：高（问卷 + 埋点 + 访谈多源收敛）"},
        {"type": "chart",
         "spec": {"type": "bar",
                  "title": "各功能满意度对比",
                  "labels": ["注册", "登录", "转账"],
                  "series": [["平均分", [2.8, 3.5, 3.9]]],
                  "figsize": [7, 4],
                  "width": 4.8,
                  "bar_width": 0.5},
         "caption": "图1 各功能满意度对比"},
        {"type": "table", "title": "参与者概况",
         "headers": ["维度", "人数", "占比"],
         "rows": [["男", 80, "51%"], ["女", 76, "49%"]]},
        {"type": "heading", "text": "四、行动建议", "level": 1},
        {"type": "bullets",
         "items": [
            "【高优先级】精简注册表单字段至必填 3 项，预期可提升注册完成率约 15%（支撑：注册流失主要集中在信息填写步骤）",
            "【中优先级】增加注册进度提示与中途保存，降低用户放弃成本（支撑：访谈 P03/P05 提及）",
            "【待验证】向新用户推送“一键登录”入口，可用 A/B 测试验证（证据不足，需补实验数据）"
         ]}
    ]
}


# --------------------------------------------------------------------------- #
# CLI 入口
# --------------------------------------------------------------------------- #
def main(argv):
    if len(argv) == 1:
        # 无参数：输出内置示例，用于自检
        out_docx = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                                "analysis_example.docx")
        data = ANALYSIS_EXAMPLE
    elif len(argv) == 2:
        # 单参数：被误当作 JSON 输入常见，给出用法提示并输出示例
        print("参数不完整：需要 <analysis.json> <output.docx>。")
        print("检测到只传入 1 个参数 '%s'，已改走无参示例模式。" % argv[1])
        out_docx = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                                "analysis_example.docx")
        data = ANALYSIS_EXAMPLE
    else:
        in_json = argv[1]
        out_docx = argv[2]
        with open(in_json, "r", encoding="utf-8") as f:
            data = json.load(f)
    build_report(data, out_docx)


if __name__ == "__main__":
    main(sys.argv)
